import csv
import json
import re
import shutil
import statistics
import sys
from functools import reduce
from math import sqrt, acos
import docx
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import scipy
import scipy.io
import scipy.signal
from PyQt5.QtWidgets import QFileDialog

matplotlib.use('Qt5Agg')
import os

from PyQt5 import QtWidgets

from PyQt5 import QtCore

import design_1_3 as design
import add_parameter as parameter_gui
import marker_rename as marker_gui
import file_load as load_gui
import pyqtgraph


class PandasModel(QtCore.QAbstractTableModel):
    def __init__(self, df=pd.DataFrame(), parent=None):
        QtCore.QAbstractTableModel.__init__(self, parent=parent)
        self._df = df

    def headerData(self, section, orientation, role=QtCore.Qt.DisplayRole):
        if role != QtCore.Qt.DisplayRole:
            return QtCore.QVariant()

        if orientation == QtCore.Qt.Horizontal:
            try:
                return self._df.columns.tolist()[section]
            except (IndexError,):
                return QtCore.QVariant()
        elif orientation == QtCore.Qt.Vertical:
            try:
                # return self.df.index.tolist()
                return self._df.index.tolist()[section]
            except (IndexError,):
                return QtCore.QVariant()

    def data(self, index, role=QtCore.Qt.DisplayRole):
        if role != QtCore.Qt.DisplayRole:
            return QtCore.QVariant()

        if not index.isValid():
            return QtCore.QVariant()

        return QtCore.QVariant(str(self._df.iloc[index.row(), index.column()]))

    def setData(self, index, value, role):
        row = self._df.index[index.row()]
        col = self._df.columns[index.column()]
        if hasattr(value, 'toPyObject'):
            # PyQt4 gets a QVariant
            value = value.toPyObject()
        else:
            # PySide gets an unicode
            dtype = self._df[col].dtype
            if dtype != object:
                value = None if value == '' else dtype.type(value)
        self._df.set_value(row, col, value)
        return True

    def rowCount(self, parent=QtCore.QModelIndex()):
        return len(self._df.index)

    def columnCount(self, parent=QtCore.QModelIndex()):
        return len(self._df.columns)

    def sort(self, column, order):
        colname = self._df.columns.tolist()[column]
        self.layoutAboutToBeChanged.emit()
        self._df.sort_values(colname, ascending=order == QtCore.Qt.AscendingOrder, inplace=True)
        self._df.reset_index(inplace=True, drop=True)
        self.layoutChanged.emit()


class addParameter(QtWidgets.QMainWindow,parameter_gui.Ui_MainWindow):
    def __init__(self,root):
        super().__init__()
        self.setupUi(self)
        self.main = root
        self.pushButton.clicked.connect(self.write)
        self.pushButton_2.clicked.connect(self.btnClosed)
    def write(self):
        name_json = 'parameters_set.json'
        name = self.lineEdit.text()
        calc = self.lineEdit_2.text()
        calc = re.sub(r'[+|-]', (lambda x: '|'+str(x.group(0))), calc)
        plane = None
        if self.radioButton_3.isChecked():
            plane = 0 # ????????????????????????
        if self.radioButton_2.isChecked():
            plane = 1 # ??????????????????????
        if self.radioButton.isChecked():
            plane = 2# ????????????????????
        if os.path.isfile(name_json):
            with open(name_json, 'r') as f:
                parameters = json.load(f)
            parameters['parameters'].append({'parameter_name': name,
                                             'parameter_calc': calc,
                                             'anatomical_plane': plane})
            with open('parameters_set.json', 'w') as f:
                json.dump(parameters, f)
        else:
            parameters = {'parameters' : [{'parameter_name': name,
                                             'parameter_calc': calc,
                                             'anatomical_plane': plane}]}
            with open('parameters_set.json', 'w') as f:
                json.dump(parameters, f)
        self.main.load_parameters()
        self.close()

    def btnClosed(self):
        self.close()
class RenameMarker(QtWidgets.QMainWindow,marker_gui.Ui_MainWindow):
    def __init__(self,data = None,path = None):
        super().__init__()
        self.setupUi(self)
        self.data = data
        self.path = path
        if self.data is None:
            pass
        else:
            work_col = [i.split()[0] for i in self.data.columns if 'X' in i]
            self.comboBox_2.addItems(work_col)
            self.pushButton.clicked.connect(self.rename)
        self.pushButton_2.clicked.connect(self.btnClosed)
    def rename(self):

        self.data[self.lineEdit.text() + ' X'] =  self.data[self.comboBox_2.currentText()+' X']
        self.data[self.lineEdit.text() + ' Y'] = self.data[self.comboBox_2.currentText() + ' Y']
        self.data[self.lineEdit.text() + ' Z'] = self.data[self.comboBox_2.currentText() + ' Z']

        drop_col = [self.comboBox_2.currentText()]
        _col = [i for i in self.data.columns if i.split()[0] in drop_col]
        self.data = self.data.drop(_col, axis=1)
        with open(self.path, 'r') as f:
            data = f.read()

        lst_ = [[j for j in i.split('\t')] for i in data.split('\n')]

        i = 0
        while lst_[i][0] != 'TRAJECTORY_TYPES':
            i += 1
        j = 0
        while lst_[j][0] != 'MARKER_NAMES':
            j += 1
        lst_[j] = (['MARKER_NAMES'] + list(set([i.split()[0] for i in self.data.columns if i != 'Frame' and i != 'Time'])))

        lst_[i + 1] = [i for i in self.data.columns]
        lst_[i + 2:] = self.data.values.tolist()

        with open(self.path, 'wt', newline='') as out_file:
            tsv_writer = csv.writer(out_file, delimiter='\t')
            for i in range(len(lst_)):
                tsv_writer.writerow(lst_[i])

    def btnClosed(self):
        self.close()

class file_load(QtWidgets.QMainWindow,load_gui.Ui_MainWindow):
    def __init__(self,root):
        super().__init__()
        self.setupUi(self)
        self.main = root
        self.pushButton_add_path.clicked.connect(self.add_path)
        self.pushButton_ok.clicked.connect(self.savepath)
        self.pushButton_close.clicked.connect(self.btnClosed)
        self.pushButton_ok_close.clicked.connect(self.save_close)

    def save_close(self):
        self.savepath()
        self.btnClosed()

    def savepath(self):
        path = self.lineEdit.text()
        name_ = self.lineEdit_2.text()
        type_ = None
        if self.firstButton.isChecked():
            type_ = 0
        if self.secondButton.isChecked():
            type_ = 1
        with open(path, 'r') as f:
            data = f.read()
        lst_ = [[j for j in i.split('\t')] for i in data.split('\n')]
        i = 0
        while lst_[i][0] != 'TRAJECTORY_TYPES':
            i += 1
        j = 0
        while lst_[j][0] != 'MARKER_NAMES':
            j += 1
        market_name = lst_[j][1:]
        if len(lst_[i + 2:][0]) - len(lst_[i + 1]) < 0:
            data = pd.DataFrame(lst_[i + 2:], columns=lst_[i + 1][:-1]).dropna()
        else:
            data = pd.DataFrame(lst_[i + 2:], columns=lst_[i + 1]).dropna()
        out ={}
        out['data'] = data
        out['market_name'] = market_name
        out['name'] = name_
        out['type'] = type_
        self.main.download_file(out)

    def add_path(self):
        filename = QFileDialog.getOpenFileName(filter='*.tsv')
        path = filename[0]
        self.lineEdit.setText(path)



    def btnClosed(self):
        self.close()

class ExampleApp(QtWidgets.QMainWindow, design.Ui_MainWindow):
    def __init__(self):
        # ?????? ?????????? ?????????? ?????? ?????????????? ?? ????????????????????, ??????????????
        # ?? ??.??. ?? ?????????? design.py
        super().__init__()
        self.setupUi(self)
        self.time_index = 0
        self.data= None
        self.name_mini_plot = None
        self.point_HS, self.point_TO = None, None
        self.parameters = None
        self.names_news = None
        self.type_of_data = None
        self.datas = []
        if os.path.isfile('parameters_set.json') == False:
            with open('parameters_set.json', 'w+') as f:
                parameters = {'parameters': [{'parameter_name': 'SVA',
                                              'parameter_calc': 'CV7|-LV5',
                                              'anatomical_plane': 0}]}
                json.dump(parameters, f)
        self.load_parameters()
        self.w = None
        self.path = None
        self.parameter_plot.setBackground('w')
        self.comboBox.currentTextChanged.connect(self.mini_plots)
        self.time_index = 1
        self.data_calc = None
        self.horizontalSlider.valueChanged.connect(self.change_value_1)
        self.pushButton.clicked.connect(self.change_value)  # ?????????????????? ????????????????????
        self.pushButton_2.clicked.connect(self.export_to_word)
        self.planeBox_s.currentTextChanged.connect(self.saggital_plot)
        self.planeBox_f.currentTextChanged.connect(self.frontal_plot)
        self.planeBox_a.currentTextChanged.connect(self.aksial_plot)
        self.action_m.triggered.connect(self.open_marker)
        self.action_p.triggered.connect(self.open_parameter)
        self.action_l.triggered.connect(self.file_load)
        self.fileBox.currentTextChanged.connect(self.change_data_name)
        # ?????????????????????????? ?????????????????? ??????????????
    def change_data_name(self,value= None):
        for i in self.datas:
            if i['name'] == value:
                self.file_change(i)

    def file_load(self):
        if self.w is None:
            self.w = file_load(self)
            self.w.show()
        else:
            self.w = None
            self.file_load()

    def open_marker(self):
        if self.w is None:
            self.w = RenameMarker(data = self.data,path = self.path)
            self.w.show()

        else:
            self.w = None
            self.open_marker()
    def open_parameter(self):
        if self.w is None:
            self.w = addParameter(self)
            self.w.show()
        else:
            self.w = None
            self.open_parameter()

    def file_change(self,out):
        self.data = out['data']
        self.market_name = out['market_name']
        work_col = [i for i in self.data.columns if i.find('Type') == -1]
        self.data = self.data[work_col].astype(float)
        self.type_of_data = out['type']
        # ?????????????????? ???????????????? ?????????????????? ?????????????? ???? ???????? ????????????
        if self.type_of_data == 1:
            self.horizontalSlider.setMaximum(self.data.shape[0])
            self.horizontalSlider.setMinimum(1)
        else:
            self.horizontalSlider.setMaximum(1)
            self.horizontalSlider.setMinimum(1)
        # ?????????????????? ?????????????? ????????
        self.showMaximized()
        self.showNormal()
        # ???????????? ????????????
        self.calc()
        self.change_value(None)
        # ?????????? ????????????????
        self.mini_plots(value=self.names_news[0])
        self.change_type()

    def download_file(self,out = None):
        if out is  None:
            pass
        else:
            self.datas.append(out)
            self.fileBox.clear()
            names = [i['name'] for i in self.datas]
            self.fileBox.addItems(names)
            self.fileBox.setCurrentText(out['name'])

    def export_to_word(self):
        # ?????????????????? ????????????????????
        if os.path.exists('./export') == False:
            os.makedirs('./export')
        doc = docx.Document()
        table = doc.add_table(rows=0, cols=3)
        lst_parameter = [i['parameter_name'] for i in self.parameters['parameters']]
        if self.type_of_data == 1:
            data_gait = self.data_gait
            count_phase = int(statistics.mean([k - j for j, k in zip(self.point_HS, self.point_TO)]))
            for i in range(len(lst_parameter)):
                lst_ = [j for j in data_gait.columns if lst_parameter[i] == j.split()[0]]
                df = data_gait[lst_]
                fig, axes = plt.subplots(figsize=(40, 40))
                x_arrange = np.arange(len(df)) / len(df) * (100 / (np.arange(len(df))[-1] / len(df)))
                ma = df.T.mean().T
                mstd = df.T.std().T
                axes.set(xlabel='???????? ???????? (%)', ylabel='???????????????? ????')
                axes.set_xlim([0, 100])
                axes.plot(x_arrange, ma, "b");
                axes.legend(loc='center right')
                axes.fill_between(x_arrange, ma - mstd, ma + mstd, color="b", alpha=0.2)

                name = './export/' + lst_parameter[i] + '.png'
                fig.savefig(name)
                fig, axes = plt.subplots(figsize=(40, 40))
                ds_1 = pd.DataFrame()
                ds_2 = pd.DataFrame()
                for j in range(len(lst_)):
                    ds_1[lst_[j] + ' ???????? ????????????????'] = data_gait[lst_[j]]
                    ds_2[lst_[j] + ' ???????? ????????????????'] = data_gait[lst_[j]][:count_phase]
                df_1 = ds_1.copy()
                df_2 = ds_2.copy()
                ma_1 = df_1.T.mean().T
                mstd_1 = df_1.T.std().T
                ma_2 = df_2.T.mean().T
                mstd_2 = df_2.T.std().T
                df_1['???????? ????????'] = np.arange(len(df_1)) / len(df_1) * (100 / (np.arange(len(df_1))[-1] / len(df_1)))
                df_2['???????? ????????'] = [k for k in
                                     np.arange(len(df_1)) / len(df_1) * (100 / (np.arange(len(df_1))[-1] / len(df_1)))][
                                    :df_2.shape[0]]
                axes.set(xlabel='???????? ???????? (%)', ylabel='???????????????? ????')
                axes.set_xlim([0, 100])
                axes.plot(df_1['???????? ????????'], ma_1, "r");
                axes.plot(df_2['???????? ????????'], ma_2, "g");
                axes.legend(loc='center right')
                axes.fill_between(df_1['???????? ????????'], ma_1 - mstd_1, ma_1 + mstd_1, color="r", alpha=0.2)
                axes.fill_between(df_2['???????? ????????'], ma_2 - mstd_2, ma_2 + mstd_2, color="g", alpha=0.2)
                name_1 = './export/' + lst_parameter[i] + ' ????????' + '.png'
                fig.savefig(name_1)
                # ?????????????????? ?????????? ?????? ??????????????
                table.style = 'Table Grid'
                row_cells = table.add_row().cells
                row_cells[0].text = lst_parameter[i]
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(name, width=1400000, height=1400000)
                paragraph = row_cells[2].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(name_1, width=1400000, height=1400000)
            stat_df = self.biomechanical_parameters(self.data_gait, lst_parameter)
        else:
            stat_df = self.biomechanical_parameters(self.data_calc, lst_parameter)
        paragraph = doc.add_paragraph('???????????????????????????? ????????????')
        paragraph.alignment = 1
        t = doc.add_table(stat_df.shape[0] + 1, stat_df.shape[1] + 1)
        # add the header rows.
        for j in range(stat_df.shape[1]):
            t.cell(0, j + 1).text = stat_df.columns[j]
        # add name of rows
        for j in range(stat_df.shape[0]):
            t.cell(j + 1, 0).text = stat_df.T.columns[j]

        # add the rest of the data frame
        for i in range(stat_df.shape[0]):
            for j in range(stat_df.shape[1]):
                t.cell(i + 1, j + 1).text = str(stat_df.values[i, j])

        doc.save(self.fileBox.currentText()+'.docx')

        shutil.rmtree('./export')


    def mini_plots(self, value):
        self.name_mini_plot = value
        df = self.data_calc
        data_y = df[self.name_mini_plot]
        data_x = [i for i in range(len(data_y))]
        data_x = [i / max(data_x) * 100 for i in data_x]
        self.parameter_plot.clear()
        self.parameter_plot.setBackground('w')
        self.parameter_plot.plot(data_x, data_y)
        self.parameter_plot.showGrid(x=True, y=True)





    def change_value_1(self, value):
        self.time_index = int(value)-1
        self.lcdNumber.display(int(value))

    def change_value(self, value):
        self.show_table()
        self.plot()

    def plot(self):
        data_dynamic = self.data  # ?????????????????? ???????????? - ?????????? ??????????????
        t = self.time_index
        lst_x = [i for i in data_dynamic.columns if i.split()[0] in self.market_name and i.split()[1] == 'X']
        lst_z = [i for i in data_dynamic.columns if i.split()[0] in self.market_name and i.split()[1] == 'Z']
        x = data_dynamic[lst_x].values[t]
        y = data_dynamic[lst_z].values[t]
        self.sagital_widget.clear()
        self.sagital_widget.setBackground('w')
        self.sagital_widget.plot(x, y, pen=None, symbol='o', symbolSize=7)
        self.sagital_widget.showGrid(x=True, y=True)
        self.sagital_widget.setRange(xRange=[min(x),max(x)])
        self.sagital_widget.setRange(yRange=[min(y),max(y)])
        lst_y = [i for i in data_dynamic.columns if i.split()[0] in self.market_name and i.split()[1] == 'Y']
        lst_z = [i for i in data_dynamic.columns if i.split()[0] in self.market_name and i.split()[1] == 'Z']
        x = data_dynamic[lst_y].values[t]
        y = data_dynamic[lst_z].values[t]
        self.front_widget.clear()
        self.front_widget.setBackground('w')
        self.front_widget.plot(x, y, pen=None, symbol='o', symbolSize=7)
        self.front_widget.showGrid(x=True, y=True)
        self.front_widget.setRange(xRange=[min(x),max(x)])
        self.front_widget.setRange(yRange=[min(y),max(y)])





    def take_data_by_name_new(self, name):
        return np.array([self.data[name + ' X'].values, self.data[name + ' Y'].values, self.data[name + ' Z'].values])

    def saggital_plot(self, value = None):
        works_col_s = [i for i in self.data_gait.columns if  i.split()[0] == value]
        df = self.data_gait[works_col_s]
        data_y = df.T.mean().T
        data_x = [i for i in range(len(data_y))]
        data_x = [i / max(data_x) * 100 for i in data_x]
        self.mean_plot_s.clear()
        self.mean_plot_s.setBackground('w')
        self.mean_plot_s.plot(data_x, data_y)
        self.mean_plot_s.showGrid(x=True, y=True)

    def frontal_plot(self, value = None):
        works_col_f = [i for i in self.data_gait.columns if  i.split()[0] == value]
        df = self.data_gait[works_col_f]
        data_y = df.T.mean().T
        data_x = [i for i in range(len(data_y))]
        data_x = [i / max(data_x) * 100 for i in data_x]
        self.mean_plot_f.clear()
        self.mean_plot_f.setBackground('w')
        self.mean_plot_f.plot(data_x, data_y)
        self.mean_plot_f.showGrid(x=True, y=True)

    def aksial_plot(self, value = None):
        works_col_a = [i for i in self.data_gait.columns if  i.split()[0] == value]
        df = self.data_gait[works_col_a]
        data_y = df.T.mean().T
        data_x = [i for i in range(len(data_y))]
        data_x = [i / max(data_x) * 100 for i in data_x]
        self.mean_plot_a.clear()
        self.mean_plot_a.setBackground('w')
        self.mean_plot_a.plot(data_x, data_y)
        self.mean_plot_a.showGrid(x=True, y=True)

    def frontal_saggital_move(self):
        if self.point_HS[0] >= self.point_TO[0]:
            self.point_TO = self.point_TO[1:]
        self.data_gait = self.calc_parameters(self.point_HS,self.point_TO)

    def biomechanical_parameters(self,data,parameters):
        df_p = None
        for parameter in parameters:
            w_ = [i for i in data.columns if i.split()[0] == parameter]
            df_ = pd.DataFrame([data[w_].max(), data[w_].min()])
            df_ = df_.rename({0: '????????????????????????', 1: '??????????????????????'})
            df_ = df_.T
            if self.type_of_data == 1:
                df_['??????????????????'] = df_.apply(lambda x: x['????????????????????????'] - x['??????????????????????'], axis=1)
                df_ = df_.T
                df_ = df_.iloc[::-1]
                df_ = df_.T
                df_ = pd.DataFrame(df_.apply(
                    lambda x: str(round(statistics.mean(x), 1)) + "+" + str(round(statistics.stdev(x), 1))),
                    columns=[parameter])
            else:
                df_ = df_.apply(lambda x: round(x))
                df_ = df_.T
            if df_p is None:
                df_p = df_
            else:
                df_p = pd.concat([df_p, df_], axis=1)
        return df_p

    def change_type(self):
        self.planeBox_a.clear()
        self.planeBox_f.clear()
        self.planeBox_s.clear()
        if self.type_of_data ==1: # ???????????????? ?????? ???????????? - ????????????????
            work_col_a = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane'] == 2]
            work_col_f = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane']==1]
            work_col_s = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane']==0]
            self.planeBox_a.addItems(work_col_a)
            self.planeBox_f.addItems(work_col_f)
            self.planeBox_s.addItems(work_col_s)

            self.planeBox_a.setCurrentIndex(0)
            self.planeBox_s.setCurrentIndex(0)
            self.planeBox_f.setCurrentIndex(0)
            df_to_model = []
            for parameters in [work_col_a,work_col_f,work_col_s]:
                stat_df = self.biomechanical_parameters(self.data_gait, parameters)
                df_to_model.append(stat_df)
            tableView_ = [self.tableView_a, self.tableView_f, self.tableView_s]
            for i in range(3):
                if df_to_model[i] is not None:
                    model_ = PandasModel(
                        df_to_model[i])
                    tableView_[i].setModel(model_)
            self.saggital_plot(self.planeBox_s.currentText())
            self.frontal_plot(self.planeBox_f.currentText())
            self.aksial_plot(self.planeBox_a.currentText())
        else:
            self.mean_plot_s.clear()
            self.mean_plot_f.clear()
            self.mean_plot_a.clear()
            self.mean_plot_a.setBackground('w')
            self.mean_plot_f.setBackground('w')
            self.mean_plot_s.setBackground('w')
            work_col_a = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane'] == 2]
            work_col_f = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane'] == 1]
            work_col_s = [i['parameter_name'] for i in self.parameters['parameters'] if i['anatomical_plane'] == 0]
            df_to_model = []
            for parameters in [work_col_a, work_col_f, work_col_s]:
                stat_df = self.biomechanical_parameters(self.data_calc, parameters)
                df_to_model.append(stat_df)
        tableView_ = [self.tableView_a, self.tableView_f, self.tableView_s]
        for i in range(3):
            if df_to_model[i] is not None:
                model_ = PandasModel(
                    df_to_model[i])
                tableView_[i].setModel(model_)

    def calc(self):
        self.data_calc = self.calc_parameters()
        self.data_calc['Time'] = self.data['Time']
        if self.type_of_data == 1:
            self.point_HS, self.point_TO = self.step_separator(self.take_data_by_name_new('R_SAE'))
            self.frontal_saggital_move()

    def load_parameters(self):
        with open('parameters_set.json', 'r') as f:
            self.parameters = json.load(f)
        self.names_news = []
        for parameter in self.parameters['parameters']:
            self.names_news.append(parameter['parameter_name'])
        self.comboBox.clear()
        self.comboBox.addItems(self.names_news)

    def calc_parameters(self,RHS=None, RTO=None):  # ???????????? ?????????????? ????????????????????
        data_calc = pd.DataFrame()
        with open('parameters_set.json', 'r') as f:
            parameters = json.load(f)
        for parameter in parameters['parameters']:
            plane_ = parameter['anatomical_plane']
            if RHS == None or RTO == None:
                t = pd.Series(np.zeros(self.data.shape[0]))
                for name_marker in parameter['parameter_calc'].split('|'):
                    function_ = (lambda x: self.take_data_by_name_new(x)[plane_])
                    if '/' in name_marker:
                        value = int(name_marker[name_marker.find('/') + 1])
                        name_marker = name_marker[:-2]
                        function_ = (lambda x, function_=function_: function_(x) / (value))
                    if '*' in name_marker:
                        value = int(name_marker[name_marker.find('*') + 1])
                        name_marker = name_marker[:-2]
                        function_ = (lambda x, function_=function_: function_(x) * (value))
                    if '-' in name_marker:
                        name_marker = name_marker[1:]
                        function_ = (lambda x, function_=function_: function_(x) * (-1))
                    if '+' in name_marker:
                        name_marker = name_marker[1:]
                        function_ = (lambda x, function_=function_: function_(x))
                    t += function_(name_marker)
                data_calc[parameter['parameter_name']] = t
            else:
                for j in range(1, len(RHS)):
                    t = pd.Series(np.zeros(RHS[j] - RHS[j - 1]))
                    for name_marker in parameter['parameter_calc'].split('|'):
                        function_ = (lambda x: self.take_data_by_name_new(x)[plane_][RHS[j - 1]:RHS[j]])
                        if '/' in name_marker:
                            value = int(name_marker[name_marker.find('/') + 1])
                            name_marker = name_marker[:-2]
                            function_ = (lambda x, function_=function_: function_(x) / (value))
                        if '*' in name_marker:
                            value = int(name_marker[name_marker.find('*') + 1])
                            name_marker = name_marker[:-2]
                            function_ = (lambda x, function_=function_: function_(x) * (value))
                        if '-' in name_marker:
                            name_marker = name_marker[1:]
                            function_ = (lambda x, function_=function_: function_(x) * (-1))
                        if '+' in name_marker:
                            name_marker = name_marker[1:]
                            function_ = (lambda x, function_=function_: function_(x))
                        t += function_(name_marker)
                    data_calc[parameter['parameter_name'] + ' ?????? ' + str(j)] = t
        return data_calc

    def step_separator(self, marker, i_=2):

        marker_x = marker[0]
        marker_y = marker[1]
        marker_z = marker[2]
        # ???????????????? ????????????????????
        if marker_x[0] < marker_x[len(marker_x) - 1]:
            marker_y = marker_y * (-1)
            marker_z = marker_z * (-1)
        # ???????????????????? ????????????
        #  ?? ???????????? 2 HZ
        acs_AP = np.append([0, 0], np.diff(marker_y.T, n=2))
        b, a = scipy.signal.butter(11, i_, 'lowpass', fs=100)
        fil_acs_AP = scipy.signal.filtfilt(b, a, acs_AP)
        binary_acs_AP = [True if i > 0 else False for i in fil_acs_AP]
        vel_vertical = np.append([0], np.diff(marker_z.T, n=1))
        # ?????????????????? ???? ???????????????????? ???? ?????????????????????????????? ?????????????? Z
        selection = 4  # ???????????????? ???????????????????? ???? ?????????????? ?????? ???????????????? ???????????? ????????????
        half_selection = int(selection / 2)
        point_HS = []
        point_TO = []
        # ???????????? ??????????????????????????????
        HS_T = None
        T_ = False
        for i in range(half_selection, len(vel_vertical) - half_selection):
            min_left = reduce(lambda x, y: x and y,
                              [True if vel_vertical[i] < j else False for j in vel_vertical[i - half_selection:i]])
            min_right = reduce(lambda x, y: x and y,
                               [True if vel_vertical[i] < j else False for j in
                                vel_vertical[i + 1:i + half_selection + 1]])
            max_left = reduce(lambda x, y: x and y,
                              [True if vel_vertical[i] > j else False for j in vel_vertical[i - half_selection:i]])
            max_right = reduce(lambda x, y: x and y,
                               [True if vel_vertical[i] > j else False for j in
                                vel_vertical[i + 1:i + half_selection + 1]])
            if ((min_left and min_right) or (max_left and max_right)) and binary_acs_AP[i]:
                HS_T = i
                T_ = True
            if not binary_acs_AP[i] and T_:
                point_HS.append(HS_T)
                T_ = False
        # ???????????? ????????????
        T_ = True
        for i in range(half_selection, len(vel_vertical) - half_selection):
            min_left = reduce(lambda x, y: x and y,
                              [True if vel_vertical[i] < j else False for j in vel_vertical[i - half_selection:i]])
            min_right = reduce(lambda x, y: x and y,
                               [True if vel_vertical[i] < j else False for j in
                                vel_vertical[i + 1:i + half_selection + 1]])
            max_left = reduce(lambda x, y: x and y,
                              [True if vel_vertical[i] > j else False for j in vel_vertical[i - half_selection:i]])
            max_right = reduce(lambda x, y: x and y,
                               [True if vel_vertical[i] > j else False for j in
                                vel_vertical[i + 1:i + half_selection + 1]])

            if ((max_left and max_right) or (min_left and min_right)) and binary_acs_AP[i] and T_ == True:
                point_TO.append(i)
                T_ = False
            if not binary_acs_AP[i]:
                T_ = True
        return [point_HS, point_TO]


    def show_table(self):
        work_col = [i for i in self.data_calc.columns if i not in 'Time']
        if self.type_of_data ==1:
            model_1 = PandasModel(
                self.data_calc[work_col].loc[self.time_index].apply(lambda x: round(x, 8)).to_frame(name="time " + str(self.time_index+1)))
        else:
            model_1 = PandasModel(
                self.data_calc[work_col].mean().apply(lambda x: round(x, 8)).to_frame(
                    name="time " + str(self.time_index+1)))

        self.tableView.setModel(model_1)
        # ???????????????????? ?????????? ???????????????????????? ?????? ?????? ?????? ???????????? ???????????????????? ??????????????


def main():
    app = QtWidgets.QApplication(sys.argv)  # ?????????? ?????????????????? QApplication
    window = ExampleApp()  # ?????????????? ???????????? ???????????? ExampleApp
    window.show()  # ???????????????????? ????????
    app.exec_()  # ?? ?????????????????? ????????????????????


if __name__ == '__main__':  # ???????? ???? ?????????????????? ???????? ????????????????, ?? ???? ??????????????????????
    main()  # ???? ?????????????????? ?????????????? main()
